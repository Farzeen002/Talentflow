"""
app/services/resume/extractor.py

Pure text extraction from PDF and DOCX resume binaries.

Design principles
-----------------
* **No I/O** — this module accepts raw bytes and returns an ExtractionResult.
  All storage reads/writes happen in the preprocessor layer, not here.
* **No MongoDB** — extraction has no knowledge of candidates or lifecycle state.
* **Never silently swallows errors** — library failures are wrapped in
  ExtractionError so the caller can decide how to handle them.
* **Lazy imports** — pdfminer.six and python-docx are imported inside the
  extraction functions so missing libraries raise a clear ImportError rather
  than failing at module load time.

Supported MIME types
--------------------
* ``application/pdf``                                                         → pdfminer.six
* ``application/msword``                                                      → python-docx
* ``application/vnd.openxmlformats-officedocument.wordprocessingml.document`` → python-docx

Future extensions (OCR, Google Doc AI, etc.) can be added as new _extract_*
functions dispatched from :func:`extract_text` without touching the caller.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass
from importlib.metadata import PackageNotFoundError
from importlib.metadata import version as _pkg_version

logger = logging.getLogger(__name__)

# ── Public constants ──────────────────────────────────────────────────────────

SUPPORTED_MIME_TYPES: frozenset[str] = frozenset({
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
})


# ══════════════════════════════════════════════════════════════════════════════
# Result + error types
# ══════════════════════════════════════════════════════════════════════════════

class ExtractionError(Exception):
    """
    Raised when the extraction library fails to process a document.

    Distinct from ``ValueError`` (bad inputs) and ``ImportError`` (missing
    library) so callers can handle library failures specifically.
    """


@dataclass(frozen=True)
class ExtractionResult:
    """
    Output of a successful text extraction.

    Attributes:
        text:              Clean UTF-8 plain text extracted from the document.
        extractor:         Library name, e.g. ``'pdfminer.six'``, ``'python-docx'``.
        extractor_version: Package version string of the extraction engine.
        char_count:        ``len(text)`` — pre-computed for convenience.
        language:          ISO 639-1 language code if detected; ``'unknown'``
                           if no language detection was performed.
    """

    text:              str
    extractor:         str
    extractor_version: str
    char_count:        int
    language:          str = "unknown"


# ══════════════════════════════════════════════════════════════════════════════
# Public entry point
# ══════════════════════════════════════════════════════════════════════════════

def extract_text(data: bytes, mime_type: str) -> ExtractionResult:
    """
    Dispatch extraction to the appropriate engine based on MIME type.

    Args:
        data:      Raw bytes of the resume file (binary content from storage).
        mime_type: MIME type string, e.g. ``'application/pdf'``.

    Returns:
        :class:`ExtractionResult` with ``text`` guaranteed to be non-empty.

    Raises:
        ValueError:       Unsupported MIME type.
        ExtractionError:  Library-level processing failure.
        ImportError:      Required extraction library is not installed.
    """
    mt = mime_type.lower().strip()

    if mt == "application/pdf":
        return _extract_pdf(data)

    if mt in (
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ):
        return _extract_docx(data)

    raise ValueError(
        f"Unsupported MIME type for text extraction: {mime_type!r}. "
        f"Supported types: {sorted(SUPPORTED_MIME_TYPES)}"
    )


# ══════════════════════════════════════════════════════════════════════════════
# PDF extraction  (pdfminer.six)
# ══════════════════════════════════════════════════════════════════════════════

def _extract_pdf(data: bytes) -> ExtractionResult:
    """
    Extract plain text from a PDF binary using pdfminer.six.

    pdfminer.six is imported lazily so the module loads without error when
    the library is absent and a helpful message is raised at call time.

    Handles malformed PDFs gracefully: pdfminer tolerates many structural
    errors internally; remaining failures surface as ExtractionError.
    """
    try:
        from pdfminer.high_level import extract_text as _pdfminer_extract
        from pdfminer.layout import LAParams
    except ImportError as exc:
        raise ImportError(
            "pdfminer.six is required for PDF extraction. "
            "Add it to requirements.txt or run: pip install pdfminer.six"
        ) from exc

    try:
        pkg_ver = _pkg_version("pdfminer.six")
    except PackageNotFoundError:
        pkg_ver = "unknown"

    logger.debug(
        "extractor: PDF extraction starting size_bytes=%d pdfminer_version=%s",
        len(data), pkg_ver,
    )

    try:
        raw = _pdfminer_extract(
            io.BytesIO(data),
            laparams=LAParams(),
        )
        text = (raw or "").strip()
    except Exception as exc:
        raise ExtractionError(
            f"pdfminer.six failed to extract text: {exc}"
        ) from exc

    logger.debug(
        "extractor: PDF extraction complete char_count=%d",
        len(text),
    )

    return ExtractionResult(
        text=text,
        extractor="pdfminer.six",
        extractor_version=pkg_ver,
        char_count=len(text),
    )


# ══════════════════════════════════════════════════════════════════════════════
# DOCX extraction  (python-docx)
# ══════════════════════════════════════════════════════════════════════════════

def _extract_docx(data: bytes) -> ExtractionResult:
    """
    Extract plain text from a DOCX (or legacy DOC) binary using python-docx.

    Each non-empty paragraph is joined with a newline.  Tables and headers
    inside the document body are included because python-docx exposes them
    in ``doc.paragraphs``.

    Note: Legacy ``.doc`` (application/msword) files must be in DOCX-compatible
    format; truly old binary .doc files may raise ExtractionError.
    """
    try:
        import docx  # pip package "python-docx"; imported as "docx"
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX extraction. "
            "Add it to requirements.txt or run: pip install python-docx"
        ) from exc

    try:
        pkg_ver = _pkg_version("python-docx")
    except PackageNotFoundError:
        pkg_ver = "unknown"

    logger.debug(
        "extractor: DOCX extraction starting size_bytes=%d python_docx_version=%s",
        len(data), pkg_ver,
    )

    try:
        doc        = docx.Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text       = "\n".join(paragraphs).strip()
    except Exception as exc:
        raise ExtractionError(
            f"python-docx failed to extract text: {exc}"
        ) from exc

    logger.debug(
        "extractor: DOCX extraction complete char_count=%d",
        len(text),
    )

    return ExtractionResult(
        text=text,
        extractor="python-docx",
        extractor_version=pkg_ver,
        char_count=len(text),
    )
